import os
from typing import List, TypedDict, Annotated
from langgraph.graph import StateGraph  # For state graph representation
from langchain_core.messages import HumanMessage, AIMessage  # To handle conversation messages
from langchain_core.prompts import ChatPromptTemplate  # To format prompts for LLMs
from langchain_core.runnables.graph import MermaidDrawMethod  # Graph drawing method
from langchain_groq import ChatGroq  # To interface with the ChatGroq LLM
from IPython.display import display  # To display outputs in Jupyter/IPython environments
import gradio as gr  # For creating a user interface
from fpdf import FPDF  # For creating PDF files
import io  # For creating downloadable files
from docx import Document  # For creating Word documents

# Define the state of the planner
class PlannerState(TypedDict):
    messages: Annotated[List[HumanMessage | AIMessage], "Messages in the conversation"]
    city: str
    interests: List[str]
    itinerary: str

# Initialize the LLM with the Groq API
llm = ChatGroq(
    temperature=0,
    groq_api_key="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",  # Replace with your valid API key
    model_name="llama-3.3-70b-versatile"
)

# Define the prompt template for generating the itinerary
itinerary_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a helpful travel assistant. Create a day trip itinerary for {city} based on the user's interests: {interests}. Provide a brief, bulleted itinerary."),
    ("human", "Create an itinerary for my day trip."),
])

# Function to input city into the state
def input_city(city: str, state: PlannerState) -> PlannerState:
    return {
        **state,
        "city": city,
        "messages": state['messages'] + [HumanMessage(content=f"I want to plan a trip to {city}.")],
    }

# Function to input interests into the state
def input_interests(interests: str, state: PlannerState) -> PlannerState:
    return {
        **state,
        "interests": interests.split(", "),
        "messages": state['messages'] + [HumanMessage(content=f"My interests are: {interests}.")],
    }

# Function to generate an itinerary using the LLM
def create_itinerary(state: PlannerState) -> str:
    response = llm.invoke(itinerary_prompt.format_messages(
        city=state['city'],
        interests=", ".join(state['interests'])
    ))
    state["itinerary"] = response.content
    state["messages"] += [AIMessage(content=response.content)]
    return response.content

# Function to create a downloadable PDF
def generate_pdf(itinerary: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, itinerary)
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.read()

# Function to create a downloadable DOC
def generate_doc(itinerary: str) -> bytes:
    doc = Document()
    doc.add_heading("Generated Itinerary", level=1)
    doc.add_paragraph(itinerary)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# Wrapper function to integrate the planning steps
def travel_planner(city: str, interests: str):
    state: PlannerState = {
        "messages": [],
        "city": "",
        "interests": [],
        "itinerary": "",
    }

    state = input_city(city, state)
    state = input_interests(interests, state)

    itinerary = create_itinerary(state)

    pdf_file = generate_pdf(itinerary)
    doc_file = generate_doc(itinerary)

    return itinerary, pdf_file, doc_file

# Gradio Interface
def interface_fn(city: str, interests: str):
    itinerary, pdf_file, doc_file = travel_planner(city, interests)
    return {
        "itinerary": itinerary,
        "pdf": (pdf_file, "itinerary.pdf"),
        "doc": (doc_file, "itinerary.docx"),
    }

interface = gr.Interface(
    fn=interface_fn,
    inputs=[
        gr.Textbox(label="Enter the city for your day trip"),
        gr.Textbox(label="Enter your interests (comma-separated)"),
    ],
    outputs=[
        gr.Textbox(label="Generated Itinerary", interactive=False),
        gr.File(label="Download Itinerary (PDF)", type="binary"),
        gr.File(label="Download Itinerary (DOC)", type="binary"),
    ],
    title="Travel Itinerary Planner",
    description="Enter a city and your interests to generate a personalized day trip itinerary. You can copy or download the itinerary as a PDF or Word document."
)

# Launch the Gradio interface
interface.launch(share=False)
